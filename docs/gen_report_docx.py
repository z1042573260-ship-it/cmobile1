# -*- coding: utf-8 -*-
"""生成《烟台基站工程情报系统 — 新人入职汇报》Word 版（.docx）"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"D:\googledownload\wangluobu_vscode\docs\新人入职汇报_基站工程情报系统.docx"

doc = Document()

# 默认正文字体（中文兼容）
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(11)

def para(text, italic=False, bold=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    return p

def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")

# ===== 标题 =====
t = doc.add_heading("烟台基站工程情报系统 — 新人入职汇报", level=0)
para("汇报对象：网络部一把手 ｜ 定位：战略价值汇报（产品经理视角）", size=10,
     color=RGBColor(0x66, 0x66, 0x66))
p = para("一句话记忆点：让网络部每一次基站建设，都变成我们提前布局的机会。",
         italic=True, bold=True, color=RGBColor(0x0F, 0x6E, 0x56))

# ===== 一、开场 =====
doc.add_heading("一、开场：一句话定位（30 秒）", level=1)
para("网络部每天有多少个基站建设商机，在政府的规划、招标、施工公示网站上悄悄出现、又悄悄溜走？")
para("我的系统，让它们一个都跑不掉。", bold=True)
para("烟台基站工程情报系统，把政府网站上每一张工程公示，自动变成一条看得懂、用得上的销售情报——"
     "销售打开地图，就知道今天该去哪个区县、找哪家企业、谈什么业务。")

# ===== 二、战略问题 =====
doc.add_heading("二、战略问题：基站商机的“黄金窗口”正在流失", level=1)
para("基站建设是一套固定的时间线：规划公示 → 招标 → 施工许可 → 主体竣工。")
para("对网络部的价值在于——宏站的最佳抢建窗口，只在“规划 / 招标 / 施工前”。"
     "一旦项目进入竣工交付，施工期抢建窗口关闭，商机基本流失。")
doc.add_heading("现状痛点与业务影响", level=2)
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Light Grid Accent 1"
tbl.rows[0].cells[0].text = "现状痛点"
tbl.rows[0].cells[1].text = "业务影响"
for a, b in [
    ("商机靠人工刷政府网站、看行业群", "滞后、漏抓，等销售知道时窗口已关"),
    ("海量公告（规划/招标/施工/招商/新闻）混在一起", "有效信号被淹没，人工无法逐一甄别"),
    ("看不到“具体位置在哪、归哪个区县”", "无法转化为可执行的销售动作"),
]:
    c = tbl.add_row().cells
    c[0].text = a
    c[1].text = b
para("这直接关乎网络部两条核心 KPI：宏站建设进度、政企业务收入。", bold=True)

# ===== 三、解决方案 =====
doc.add_heading("三、解决方案：基站工程情报系统", level=1)
para("产品定位：一套“采集—研判—呈现”自动化的基站工程情报系统，让政府公示与销售动作之间，不再有信息差。")
para("技术主线（一笔带过，不展开代码）：", bold=True)
para("政府/招标网站 (7大源)  →  AI 研判管线  →  结构化数据库  →  高德地图看板")
para("   爬虫自动采集          分类/预警/商机        MySQL           红黄预警·区县下钻")
bullets([
    "采集层：自动抓取烟台市及各区县自然资源、公共资源交易、投资促进等 7 大官方源",
    "研判层：AI 管线逐条分析——判定项目本质、预警等级（红/黄）、基站需求、商机信号、精确经纬度",
    "呈现层：高德地图看板，红色预警直接落点，支持区县筛选、项目下钻、详情查看",
])

# ===== 四、三大亮点 =====
doc.add_heading("四、核心能力：三大亮点", level=1)
doc.add_heading("亮点 1 · 全自动，零人工盯屏", level=2)
bullets([
    "7 大官方源自动采集，每周稳定新增约 20 条待研判公告",
    "AI 自动完成分类、预警、商机判断、经纬度补全，不再依赖人工逐条刷网站",
])
doc.add_heading("亮点 2 · 看得懂的地图", level=2)
bullets([
    "红色预警直接落在高德地图上，销售打开即知“哪个区县、什么项目、具体位置”",
    "支持按区县筛选、按预警等级下钻，指挥调度一目了然",
])
doc.add_heading("亮点 3 · 抢时间，赢在窗口前", level=2)
bullets([
    "规划公示当天即完成研判推送，比竞争对手早数天甚至数周切入",
    "宏站抢建窗口期内精准触达，把“建好再找”变成“建前就谈”",
])

# ===== 五、应用成果 =====
doc.add_heading("五、应用成果：数据说话", level=1)
tbl2 = doc.add_table(rows=1, cols=2)
tbl2.style = "Light Grid Accent 1"
tbl2.rows[0].cells[0].text = "指标"
tbl2.rows[0].cells[1].text = "成果"
for a, b in [
    ("累计研判公告", "702 条原始 → 144 条有效红/黄预警"),
    ("区县覆盖", "烟台全域 14 个区县全覆盖"),
    ("预警结构", "红色（高优先）89 条 / 黄色 55 条"),
    ("更新频率", "每周稳定增量更新，常态运行"),
    ("典型项目", "工商学院龙口校区（9 栋楼施工总承包）、烟台南站枢纽、威思顿产业园等"),
]:
    c = tbl2.add_row().cells
    c[0].text = a
    c[1].text = b
para("视觉主证据：高德看板红点分布图（覆盖 14 区县的商机热力），是整场汇报冲击最强的一页。",
     italic=True)

# ===== 六、未来规划 =====
doc.add_heading("六、未来规划", level=1)
bullets([
    "自动定时分析：每日/每周自动跑全链，生成商机日报，无需人工触发",
    "多模型兜底：主模型 + 免费模型自动切换，保障高峰期稳定运行",
    "商机跟进闭环：预警 → 派单 → 跟进 → 成单，形成可度量的转化链路",
    "可复制推广：从烟台试点，向全省/全国基站建设情报场景复制",
])

# ===== 七、总结 =====
doc.add_heading("七、总结", level=1)
para("让网络部每一次基站建设，都变成我们提前布局的机会。", bold=True,
     color=RGBColor(0x27, 0x50, 0x0A), size=13)
para("这不是一个“爬虫工具”，而是一套把政府公开信息转化为网络部核心业绩的情报系统——"
     "它帮网络部，在每一个黄金窗口里，都快人一步。")

# ===== 附：PPT 逐页映射 =====
doc.add_page_break()
doc.add_heading("附：PPT 逐页映射", level=1)
tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = "Light Grid Accent 1"
h = tbl3.rows[0].cells
h[0].text = "PPT 页"; h[1].text = "对应章节"; h[2].text = "视觉主角"
for a, b, c in [
    ("P1 封面", "标题 + 一句话定位", "项目名 + 记忆点"),
    ("P2 钩子", "开场 30 秒", "一句反问 + 一句承诺"),
    ("P3 战略问题", "第二章", "KPI 绑定图（黄金窗口时间线）"),
    ("P4 解决方案", "第三章", "技术主线架构图（采集→研判→呈现）"),
    ("P5-7 三大亮点", "第四章", "每页一个 wow 点 + 截图"),
    ("P8 应用成果", "第五章", "高德看板红点图 + 数据卡"),
    ("P9 未来规划", "第六章", "路线图"),
    ("P10 总结", "第七章", "记忆点大字"),
]:
    cells = tbl3.add_row().cells
    cells[0].text = a
    cells[1].text = b
    cells[2].text = c

doc.save(OUT)
print("已生成:", OUT)
